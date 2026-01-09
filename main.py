"""
Live Translation Overlay (EN -> Target Language) with Large-Scale History (Word Export, Auto-Split)

Key features
- Real-time microphone capture (sounddevice)
- English ASR via faster-whisper (WhisperModel)
- Audio overlap + text de-duplication to reduce boundary drops/repeats
- Sentence splitting + idle finalize
- Translation via NLLB-200 distilled 1.3B (Transformers), target language selectable at runtime
- Two-window UI: semi-transparent background + crisp chroma-key overlay subtitles
- Persistent history stored in SQLite (async, best-effort; may drop records under overload; limited by disk)
- Export full history of the CURRENT SESSION to Word (.docx), auto-splitting into multiple files (by record count)

Threading model
- Main thread: Tkinter UI
- ASR worker thread: transcribe + translate + commit
- Logger thread: async SQLite inserts (disk IO won't stall subtitles)

NOTES
- The Target Language dropdown shows ENGLISH NAMES (not NLLB codes).
- For best language naming, install: pip install pycountry
- First run downloads models from Hugging Face (Whisper/NLLB). Internet + disk cache required.
- Chroma-key transparency via Tk "-transparentcolor" is generally Windows-only; on other OS the key color may be visible.

Open-source friendly changes in this version
- No hard-coded HF_HOME. Use CLI --hf-home or env HF_HOME if you want.
- Model load is deferred until main() so CLI/env overrides take effect.
"""

from __future__ import annotations

import argparse
import time
import re
import os
import warnings
import threading
import queue
import ctypes  # High DPI
import platform
import uuid
import sqlite3
from typing import Optional, List, Tuple

import torch
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from transformers.utils import logging as hf_logging

import tkinter as tk
from tkinter import colorchooser, filedialog, messagebox
from tkinter import ttk

import numpy as np

# Optional: pycountry for language English names (recommended)
try:
    import pycountry  # pip install pycountry
    HAS_PYCOUNTRY = True
except Exception:
    pycountry = None
    HAS_PYCOUNTRY = False

# Optional: faster-whisper (ASR)
try:
    from faster_whisper import WhisperModel
    HAS_WHISPER = True
except Exception:
    WhisperModel = None
    HAS_WHISPER = False

# Optional: sounddevice (audio capture)
try:
    import sounddevice as sd
    HAS_SD = True
except Exception:
    sd = None
    HAS_SD = False


def enable_dpi_awareness():
    """Improve UI scaling on Windows high-DPI displays. No-op on non-Windows."""
    try:
        ctypes.windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        try:
            ctypes.windll.user32.SetProcessDPIAware()
        except Exception:
            pass


enable_dpi_awareness()

# ===================== Base Configuration =====================

os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# IMPORTANT (GitHub-friendly):
# Do NOT hard-code HF_HOME here. Let Hugging Face use its default cache location,
# or let users override via:
#   - env: HF_HOME=/path/to/cache
#   - CLI: --hf-home /path/to/cache
#
# If you insist on giving a default, prefer a cross-platform one (commented out):
# os.environ.setdefault("HF_HOME", os.path.join(os.path.expanduser("~"), ".cache", "huggingface"))

# Enable hf_transfer by default if user didn't set it (can be disabled via CLI)
os.environ.setdefault("HF_HUB_ENABLE_HF_TRANSFER", "1")

hf_logging.set_verbosity_error()
warnings.filterwarnings("ignore", message="Recommended: pip install sacremoses.")

# --- MT (Translation) ---
NLLB_NAME = "facebook/nllb-200-distilled-1.3B"

# --- ASR (Whisper) ---
WHISPER_MODEL_NAME = "large-v3"
AUDIO_SAMPLE_RATE = 16000

# ===================== Tuning Knobs (Latency vs Accuracy) =====================

POLL_INTERVAL = 0.25
WHISPER_MIN_SECONDS = 1.4
WHISPER_CHUNK_SECONDS = 4.2
AUDIO_OVERLAP_SECONDS = 0.6  # recommended range ~0.4-0.8s
IDLE_FINALIZE_SECONDS = 2.8
MIN_FINALIZE_CHARS = 18

# UI ring buffer size (kept small to guarantee smooth rendering)
MAX_HISTORY_SEGMENTS = 200

# Text de-duplication overlap window
MAX_OVERLAP_CHARS = 120

# Word export split size: max records per .docx file (avoids huge .docx files).
# NOTE: splitting is by number of persisted records, NOT by "lines" or character count.
DOCX_MAX_RECORDS_PER_FILE = 2000

# SQLite storage location (override via env or CLI if desired).
DEFAULT_DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "translation_history.sqlite3")
HISTORY_DB_PATH = None  # set in main() after env/CLI apply

# Session ID: export uses this to export "this run only".
SESSION_ID = f"{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"

# Default translation target (NLLB language code token)
DEFAULT_TARGET_LANG = "zho_Hans"

# Chroma-key color for overlay background (made transparent on Windows via -transparentcolor)
KEY_COLOR = "#010203"

WHISPER_CT2_ALIASES = {
    "large-v3-turbo": [
        "h2oai/faster-whisper-large-v3-turbo",
        "Infomaniak-AI/faster-whisper-large-v3-turbo",
        "deepdml/faster-whisper-large-v3-turbo-ct2",
        "freddierice/openwhisper-turbo-large-v3-ct2",
    ],
}


def resolve_whisper_model_id(name: str):
    """
    Resolve model name into a list of candidates.
    Priority:
      1) WHISPER_CT2_REPO env var
      2) Local folder
      3) Aliases
      4) Direct identifier
    """
    if not name:
        return []
    env_repo = os.environ.get("WHISPER_CT2_REPO", "").strip()
    if env_repo:
        return [env_repo]
    if os.path.isdir(name):
        return [name]
    if name in WHISPER_CT2_ALIASES:
        return WHISPER_CT2_ALIASES[name]
    return [name]


# ===================== Global State =====================

tok = None
mt_model = None
MT_CFG = None  # {"device","forced_bos","dtype","target_lang"}  target_lang stores NLLB CODE

# Protect MT_CFG + model.generate usage, since forced_bos can change at runtime
mt_lock = threading.Lock()

whisper_lock = threading.Lock()
whisper_model = None
whisper_loaded = False

audio_stream = None
audio_q = queue.Queue(maxsize=80)
AUDIO_DEVICE_INDEX = None
audio_buffer = np.zeros(0, dtype=np.float32)

pending_en = ""
current_en = ""
current_out = ""  # translated output for current line
history_segments: List[Tuple[str, str]] = []  # UI ring buffer: (en, out_text)
asr_last_text_ts = 0.0

RUNNING = True

# Filled after tokenizer load
NLLB_ALL_LANG_CODES: List[str] = []

# UI display mapping (ENGLISH LABELS)
NLLB_LANG_LABELS: List[str] = []
NLLB_LABEL_TO_CODE = {}
NLLB_CODE_TO_LABEL = {}


# ===================== MT Loading =====================

def _pick_mt_device():
    return torch.device("cuda") if torch.cuda.is_available() else torch.device("cpu")


def _resolve_lang_id(lang_code: str) -> Optional[int]:
    """
    Resolve NLLB language code token id using tokenizer first, then model config maps.
    Returns None if not resolvable.
    """
    global tok, mt_model
    code = (lang_code or "").strip()
    if not code or tok is None or mt_model is None:
        return None

    try:
        tid = tok.convert_tokens_to_ids(code)
        if tid is not None and tid != tok.unk_token_id:
            return int(tid)
    except Exception:
        pass

    # Fallback: some builds expose maps in model config
    for attr in ["lang_code_to_id", "target_lang2id", "lang2id"]:
        try:
            if hasattr(mt_model.config, attr):
                m = getattr(mt_model.config, attr)
                if isinstance(m, dict) and code in m:
                    return int(m[code])
        except Exception:
            pass

    return None


def set_target_language(lang_code: str) -> bool:
    """
    Update translation target language at runtime.
    This does NOT reload the model; it only updates forced_bos_token_id.
    """
    global MT_CFG
    lang_code = (lang_code or "").strip()
    if not lang_code or MT_CFG is None:
        return False

    tid = _resolve_lang_id(lang_code)
    if tid is None:
        return False

    with mt_lock:
        MT_CFG["forced_bos"] = tid
        MT_CFG["target_lang"] = lang_code
    return True


def load_mt_model():
    """
    Loads the NLLB translation model/tokenizer.
    NOTE: First run will download model weights from Hugging Face if not cached.
    """
    device = _pick_mt_device()
    dtype = torch.float16 if device.type == "cuda" else torch.float32
    try:
        print(f"[MT] Loading: {NLLB_NAME} on {device.type.upper()} dtype={dtype}")
        tok_local = AutoTokenizer.from_pretrained(NLLB_NAME, src_lang="eng_Latn")
        model_local = AutoModelForSeq2SeqLM.from_pretrained(
            NLLB_NAME,
            torch_dtype=dtype,
            device_map=None,
        )
        model_local = model_local.to(device)
        model_local.eval()

        # Initialize to default target language
        forced_bos = tok_local.convert_tokens_to_ids(DEFAULT_TARGET_LANG)
        if forced_bos is None or forced_bos == tok_local.unk_token_id:
            forced_bos = None
            for attr in ["lang_code_to_id", "target_lang2id", "lang2id"]:
                if hasattr(model_local.config, attr):
                    m = getattr(model_local.config, attr)
                    if isinstance(m, dict) and DEFAULT_TARGET_LANG in m:
                        forced_bos = m[DEFAULT_TARGET_LANG]
                        break
        if forced_bos is None:
            raise RuntimeError(f"Cannot resolve default target language id: {DEFAULT_TARGET_LANG}")

        cfg = {"device": device, "forced_bos": int(forced_bos), "dtype": dtype, "target_lang": DEFAULT_TARGET_LANG}
        print(f"[MT] Using NLLB-1.3B ({device.type.upper()})")
        return tok_local, model_local, cfg
    except Exception as e:
        print(f"[MT] Failed to load MT model: {e}")
        raise SystemExit(1) from e


# ===================== Build NLLB language codes (robust) =====================

LANG_CODE_RE = re.compile(r"^[a-z]{3}_[A-Za-z]{4}$")  # e.g., eng_Latn, zho_Hans

def build_nllb_lang_codes(tok_obj, model_obj) -> List[str]:
    """
    Collect available NLLB language code tokens.
    This is best-effort across tokenizer/model variants.
    """
    codes = set()

    # tokenizer.lang_code_to_id (if available)
    try:
        m = getattr(tok_obj, "lang_code_to_id", None)
        if isinstance(m, dict) and m:
            codes.update(m.keys())
    except Exception:
        pass

    # tokenizer.additional_special_tokens is often where NLLB puts language tokens
    try:
        for t in (getattr(tok_obj, "additional_special_tokens", None) or []):
            if isinstance(t, str) and LANG_CODE_RE.match(t):
                codes.add(t)
    except Exception:
        pass

    # model config maps (fallback)
    for attr in ("lang_code_to_id", "target_lang2id", "lang2id"):
        try:
            m2 = getattr(getattr(model_obj, "config", None), attr, None)
            if isinstance(m2, dict) and m2:
                codes.update(m2.keys())
        except Exception:
            pass

    codes.add(DEFAULT_TARGET_LANG)
    return sorted(codes)


# ===================== Build ENGLISH labels for dropdown (no codes shown) =====================

SCRIPT_NAME = {
    "Latn": "Latin",
    "Cyrl": "Cyrillic",
    "Arab": "Arabic",
    "Deva": "Devanagari",
    "Hans": "Simplified",
    "Hant": "Traditional",
    "Hang": "Hangul",
    "Jpan": "Japanese",
    "Kore": "Korean",
    "Grek": "Greek",
    "Hebr": "Hebrew",
    "Thai": "Thai",
    "Beng": "Bengali",
    "Taml": "Tamil",
    "Telu": "Telugu",
    "Mlym": "Malayalam",
    "Knda": "Kannada",
    "Guru": "Gurmukhi",
    "Gujr": "Gujarati",
    "Orya": "Odia",
    "Sinh": "Sinhala",
    "Mymr": "Myanmar",
    "Khmr": "Khmer",
    "Laoo": "Lao",
    "Geor": "Georgian",
    "Armn": "Armenian",
    "Ethi": "Ethiopic",
}

def _iso3_to_english_name(iso3: str) -> str:
    """
    Convert ISO639-3 code to an English name.
    Uses pycountry if available; otherwise returns a generic placeholder.
    (No codes are shown in UI when pycountry is missing.)
    """
    iso3 = (iso3 or "").strip().lower()
    if not iso3:
        return "Unknown language"

    if HAS_PYCOUNTRY:
        try:
            lang = pycountry.languages.get(alpha_3=iso3)
            if not lang:
                lang = pycountry.languages.get(bibliographic=iso3) or pycountry.languages.get(terminology=iso3)
            if lang and getattr(lang, "name", None):
                return str(lang.name)
        except Exception:
            pass

    return "Unknown language"

def nllb_code_to_english_label(code: str) -> str:
    """
    Convert NLLB code like 'zho_Hans' into an English label like 'Chinese (Simplified)'.
    Does NOT show the code.
    """
    code = (code or "").strip()
    if "_" not in code:
        return _iso3_to_english_name(code)

    iso3, script = code.split("_", 1)
    lang_name = _iso3_to_english_name(iso3)

    script = (script or "").strip()
    if not script:
        return lang_name

    if script == "Hans":
        return f"{lang_name} (Simplified)"
    if script == "Hant":
        return f"{lang_name} (Traditional)"

    script_label = SCRIPT_NAME.get(script, script)
    return f"{lang_name} ({script_label})"

def build_nllb_label_tables(codes: List[str]):
    """
    Build (labels, label->code, code->label) tables for the UI dropdown.
    Deduplicates labels without exposing codes; duplicates get "Option N".
    """
    items = []
    for c in (codes or []):
        lbl = nllb_code_to_english_label(c)
        items.append((lbl.lower(), lbl, c))
    items.sort(key=lambda x: x[0])

    labels = []
    label_to_code = {}
    code_to_label = {}

    seen = {}
    for _key, base_lbl, code in items:
        n = seen.get(base_lbl, 0) + 1
        seen[base_lbl] = n
        lbl = base_lbl if n == 1 else f"{base_lbl} (Option {n})"
        labels.append(lbl)
        label_to_code[lbl] = code
        if code not in code_to_label:
            code_to_label[code] = lbl

    return labels, label_to_code, code_to_label


def bootstrap_mt_and_languages():
    """Defer MT load until after CLI/env overrides."""
    global tok, mt_model, MT_CFG
    global NLLB_ALL_LANG_CODES, NLLB_LANG_LABELS, NLLB_LABEL_TO_CODE, NLLB_CODE_TO_LABEL

    tok, mt_model, MT_CFG = load_mt_model()

    NLLB_ALL_LANG_CODES = build_nllb_lang_codes(tok, mt_model)
    print(f"[MT] NLLB language codes loaded: {len(NLLB_ALL_LANG_CODES)}")

    NLLB_LANG_LABELS, NLLB_LABEL_TO_CODE, NLLB_CODE_TO_LABEL = build_nllb_label_tables(NLLB_ALL_LANG_CODES)
    print(f"[MT] NLLB language labels loaded: {len(NLLB_LANG_LABELS)} (pycountry={'ON' if HAS_PYCOUNTRY else 'OFF'})")


# ===================== Whisper Loading =====================

def load_whisper_model() -> bool:
    """
    Loads the ASR model (faster-whisper).
    NOTE: First run will download model weights if not cached.
    """
    global whisper_model, whisper_loaded
    if not HAS_WHISPER:
        print("[ASR] faster-whisper is not installed: pip install faster-whisper")
        return False

    device = "cuda" if torch.cuda.is_available() else "cpu"

    if device == "cuda":
        compute_type_candidates = ["float16"]  # force float16 on CUDA
    else:
        compute_type_candidates = ["int8", "float32"]

    model_id_candidates = resolve_whisper_model_id(WHISPER_MODEL_NAME)

    with whisper_lock:
        if whisper_loaded and whisper_model is not None:
            return True

        last_err = None
        for model_id in model_id_candidates:
            for compute_type in compute_type_candidates:
                try:
                    print(f"[ASR] Loading Whisper: {model_id} on {device}, compute_type={compute_type}")
                    whisper_model = WhisperModel(model_id, device=device, compute_type=compute_type)
                    whisper_loaded = True
                    print("[ASR] Whisper loaded.")
                    return True
                except Exception as e:
                    last_err = e
                    whisper_model = None
                    whisper_loaded = False
                    if device == "cuda":
                        try:
                            torch.cuda.empty_cache()
                        except Exception:
                            pass
                    print(f"[ASR] Whisper load failed (model={model_id}, compute_type={compute_type}): {e}")

        print(f"[ASR] Whisper load failed (all candidates). Last error: {last_err}")
        return False


# ===================== Persistent History (SQLite) =====================

LOG_Q = queue.Queue(maxsize=2000)
LOGGER_THREAD: Optional[threading.Thread] = None
LOGGER_STOP = threading.Event()


def _db_connect(path: str) -> sqlite3.Connection:
    """Create a SQLite connection with pragmatic settings for frequent inserts."""
    conn = sqlite3.connect(path, timeout=30)
    conn.execute("PRAGMA journal_mode=WAL;")
    conn.execute("PRAGMA synchronous=NORMAL;")
    conn.execute("PRAGMA temp_store=MEMORY;")
    return conn


def init_history_db(path: str) -> None:
    """
    Initialize DB schema if not present.
    Also performs a minimal migration to add missing columns for older DB versions.
    """
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    conn = _db_connect(path)
    try:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS segments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id TEXT NOT NULL,
                ts REAL NOT NULL,
                ts_str TEXT NOT NULL,
                tgt_lang TEXT NOT NULL DEFAULT '',
                en TEXT NOT NULL,
                out_text TEXT NOT NULL
            );
            """
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_segments_session_id ON segments(session_id);")

        cols = [r[1] for r in conn.execute("PRAGMA table_info(segments);").fetchall()]
        if "tgt_lang" not in cols:
            conn.execute("ALTER TABLE segments ADD COLUMN tgt_lang TEXT NOT NULL DEFAULT '';")
        if "out_text" not in cols:
            conn.execute("ALTER TABLE segments ADD COLUMN out_text TEXT NOT NULL DEFAULT '';")

        conn.commit()
    finally:
        conn.close()


def enqueue_history_record(en: str, out_text: str) -> None:
    """
    Enqueue a record to be persisted (SQLite) via the async logger thread.

    Behavior:
    - Best-effort, non-blocking
    - If queue is full, drop the OLDEST queued record
    """
    global MT_CFG
    if not en and not out_text:
        return

    with mt_lock:
        tgt = (MT_CFG or {}).get("target_lang", DEFAULT_TARGET_LANG)

    rec = (SESSION_ID, time.time(), time.strftime("%Y-%m-%d %H:%M:%S"), tgt, (en or ""), (out_text or ""))

    try:
        if LOG_Q.full():
            try:
                LOG_Q.get_nowait()
            except Exception:
                pass
        LOG_Q.put_nowait(rec)
    except Exception:
        pass


def logger_worker(db_path: str):
    """Async logger thread: batch inserts into SQLite to avoid UI/ASR stalls."""
    conn = _db_connect(db_path)
    try:
        buf = []
        last_flush = time.time()

        while not LOGGER_STOP.is_set() or not LOG_Q.empty():
            try:
                item = LOG_Q.get(timeout=0.2)
                buf.append(item)
            except queue.Empty:
                item = None

            now = time.time()
            should_flush = (len(buf) >= 50) or (buf and (now - last_flush) >= 0.5)

            if should_flush:
                try:
                    conn.executemany(
                        "INSERT INTO segments(session_id, ts, ts_str, tgt_lang, en, out_text) VALUES (?, ?, ?, ?, ?, ?);",
                        buf,
                    )
                    conn.commit()
                except Exception:
                    try:
                        conn.rollback()
                    except Exception:
                        pass
                buf.clear()
                last_flush = now

        if buf:
            try:
                conn.executemany(
                    "INSERT INTO segments(session_id, ts, ts_str, tgt_lang, en, out_text) VALUES (?, ?, ?, ?, ?, ?);",
                    buf,
                )
                conn.commit()
            except Exception:
                try:
                    conn.rollback()
                except Exception:
                    pass
    finally:
        try:
            conn.close()
        except Exception:
            pass


def start_logger():
    assert HISTORY_DB_PATH is not None
    init_history_db(HISTORY_DB_PATH)
    LOGGER_STOP.clear()
    global LOGGER_THREAD
    LOGGER_THREAD = threading.Thread(target=logger_worker, args=(HISTORY_DB_PATH,), daemon=True)
    LOGGER_THREAD.start()


def stop_logger():
    LOGGER_STOP.set()
    if LOGGER_THREAD is not None:
        try:
            LOGGER_THREAD.join(timeout=2.0)
        except Exception:
            pass


def _fetch_session_bounds(conn: sqlite3.Connection, session_id: str) -> Tuple[int, int]:
    """Return (count, max_id) for a stable export snapshot."""
    row = conn.execute(
        "SELECT COUNT(*), COALESCE(MAX(id), 0) FROM segments WHERE session_id=?;",
        (session_id,),
    ).fetchone()
    count = int(row[0]) if row else 0
    max_id = int(row[1]) if row else 0
    return count, max_id


def _fetch_session_lang_summary(conn: sqlite3.Connection, session_id: str) -> str:
    """Build a short session language summary for Word header."""
    rows = conn.execute(
        "SELECT DISTINCT tgt_lang FROM segments WHERE session_id=? ORDER BY tgt_lang ASC;",
        (session_id,),
    ).fetchall()
    langs = [str(r[0] or "").strip() for r in rows if str(r[0] or "").strip()]
    if not langs:
        return DEFAULT_TARGET_LANG
    if len(langs) == 1:
        return langs[0]
    if len(langs) <= 8:
        return "Multiple: " + ", ".join(langs)
    return f"Multiple ({len(langs)} codes)"


def iter_session_records(conn: sqlite3.Connection, session_id: str, max_id: int, batch_size: int = 500):
    """Stream records for a session up to max_id in ascending id order."""
    last_id = 0
    while True:
        rows = conn.execute(
            """
            SELECT id, ts_str, tgt_lang, en, out_text
            FROM segments
            WHERE session_id=? AND id>? AND id<=?
            ORDER BY id ASC
            LIMIT ?;
            """,
            (session_id, last_id, max_id, batch_size),
        ).fetchall()
        if not rows:
            break
        for r in rows:
            last_id = int(r[0])
            yield r  # (id, ts_str, tgt_lang, en, out_text)


# ===================== Audio Device Utilities =====================

def list_input_devices():
    devices = []
    if not HAS_SD:
        return devices
    try:
        devs = sd.query_devices()
        for i, d in enumerate(devs):
            if d.get("max_input_channels", 0) > 0:
                devices.append((i, d.get("name", f"Input {i}")))
    except Exception:
        pass
    return devices


def set_audio_device(index):
    global AUDIO_DEVICE_INDEX, audio_buffer
    AUDIO_DEVICE_INDEX = index
    try:
        stop_audio_stream()
    except Exception:
        pass
    audio_buffer = np.zeros(0, dtype=np.float32)
    ok = start_audio_stream()
    if not ok:
        print("[ASR] Failed to restart audio stream with the new device.")
    return ok


# ===================== Audio Input =====================

def _audio_callback(indata, frames, time_info, status):
    if status:
        pass
    try:
        x = indata[:, 0].copy()
        if audio_q.full():
            try:
                audio_q.get_nowait()
            except Exception:
                pass
        audio_q.put_nowait(x)
    except Exception:
        pass


def stop_audio_stream():
    global audio_stream
    if audio_stream is not None:
        try:
            audio_stream.stop()
            audio_stream.close()
        except Exception as e:
            print("[ASR] stop stream error:", e)
        audio_stream = None
    try:
        while True:
            audio_q.get_nowait()
    except queue.Empty:
        pass


def start_audio_stream() -> bool:
    global audio_stream, AUDIO_DEVICE_INDEX
    if not HAS_SD:
        print("[ASR] sounddevice is not installed: pip install sounddevice")
        return False
    if audio_stream is not None:
        return True

    try:
        device_kwargs = {}
        if AUDIO_DEVICE_INDEX is not None:
            device_kwargs["device"] = AUDIO_DEVICE_INDEX

        audio_stream = sd.InputStream(
            samplerate=AUDIO_SAMPLE_RATE,
            channels=1,
            dtype="float32",
            callback=_audio_callback,
            **device_kwargs,
        )
        audio_stream.start()
        print("[ASR] Audio stream started.")
        return True
    except Exception as e:
        print(f"[ASR] Failed to start audio stream: {e}")
        audio_stream = None
        return False


# ===================== Text Processing =====================

def normalize_audio(x: np.ndarray) -> np.ndarray:
    if x is None or x.size == 0:
        return x
    x = x.astype(np.float32, copy=False)
    m = np.max(np.abs(x))
    if m > 1.0:
        x = x / m
    return x


def is_valid_en_segment(text: str) -> bool:
    if not text:
        return False
    text = text.strip()
    if not text:
        return False
    alnum = re.findall(r"[A-Za-z0-9]", text)
    return len(alnum) >= 2


def sanitize_output_text(s: str) -> str:
    """Post-process translation text. Keep it readable and stable."""
    s = (s or "").strip()
    if not s:
        return ""
    return re.sub(r"\s+", " ", s).strip()


def split_complete_sentences(buf: str):
    buf = (buf or "").strip()
    if not buf:
        return [], ""

    parts = re.split(r"([.!?\u3002\u2026]+)", buf)
    sentences = []
    cur = ""
    for i in range(0, len(parts), 2):
        seg = parts[i].strip()
        punct = parts[i + 1] if i + 1 < len(parts) else ""
        if seg:
            s = (seg + punct).strip()
            if punct:
                sentences.append(s)
            else:
                cur = s
        else:
            if punct and cur:
                sentences.append((cur + punct).strip())
                cur = ""
    return sentences, cur


# ---- Overlap / de-duplication helpers ----

WORD_RE = re.compile(r"[A-Za-z0-9']+")


def words_with_spans(s: str):
    return [(m.group(0), m.start(), m.end()) for m in WORD_RE.finditer(s or "")]


def word_overlap_suffix_prefix(a: str, b: str, max_words: int = 8) -> int:
    if not a or not b:
        return 0
    wa = [w.lower() for w, _, _ in words_with_spans(a)]
    wb = [w.lower() for w, _, _ in words_with_spans(b)]
    if not wa or not wb:
        return 0
    max_len = min(max_words, len(wa), len(wb))
    for k in range(max_len, 0, -1):
        if wa[-k:] == wb[:k]:
            return k
    return 0


def drop_prefix_words(text: str, k_words: int) -> str:
    if k_words <= 0:
        return text or ""
    spans = words_with_spans(text or "")
    if len(spans) < k_words:
        return text or ""
    cut = spans[k_words - 1][2]
    return (text[cut:] if text else "").lstrip(" \t\r\n,.;:!?")


def overlap_suffix_prefix(a: str, b: str, max_k: int) -> int:
    if not a or not b:
        return 0
    a = a[-max_k:]
    max_len = min(len(a), len(b))
    for k in range(max_len, 0, -1):
        if a[-k:].lower() == b[:k].lower():
            return k
    return 0


def append_with_overlap(existing: str, new_text: str, max_k: int = MAX_OVERLAP_CHARS) -> str:
    existing = existing or ""
    new_text = new_text or ""
    if not existing:
        return new_text.strip()
    if not new_text:
        return existing.strip()

    k_words = word_overlap_suffix_prefix(existing, new_text, max_words=8)
    if k_words >= 2:
        trimmed = drop_prefix_words(new_text, k_words)
        merged = (existing.rstrip() + " " + trimmed.lstrip()).strip()
        return re.sub(r"\s+", " ", merged)

    k = overlap_suffix_prefix(existing, new_text, max_k=max_k)
    if k > 0:
        merged = (existing + new_text[k:]).strip()
        return re.sub(r"\s+", " ", merged)

    ex = existing.strip()
    nw = new_text.strip()
    if ex and nw and ex[-1].isalnum() and nw[0].isalnum():
        merged = (ex + " " + nw).strip()
    else:
        merged = (ex + " " + nw).strip()
    return re.sub(r"\s+", " ", merged)


# ===================== Translation =====================

def translate_en_to_target(text_en: str) -> str:
    """
    Translate English text into the currently selected NLLB target language.
    Target language is controlled by MT_CFG["forced_bos"] and MT_CFG["target_lang"] (code).
    """
    global tok, mt_model, MT_CFG
    text_en = (text_en or "").strip()
    if not text_en or tok is None or mt_model is None or MT_CFG is None:
        return ""

    with mt_lock:
        device = MT_CFG["device"]
        forced_bos = MT_CFG["forced_bos"]

    try:
        inputs = tok(text_en, return_tensors="pt", truncation=True, max_length=256)
        inputs = {k: v.to(device) for k, v in inputs.items()}

        with mt_lock:
            out = mt_model.generate(
                **inputs,
                max_new_tokens=192,
                num_beams=4,
                do_sample=False,
                forced_bos_token_id=forced_bos,
                repetition_penalty=1.05,
            )

        out_text = tok.batch_decode(out, skip_special_tokens=True)[0]
        return sanitize_output_text(out_text)
    except Exception as e:
        print(f"[MT] translate error: {e}")
        return ""


# ===================== UI (Two Windows: Transparent Background + Crisp Subtitles) =====================

class OverlaySubtitleApp:
    """
    bg_root: semi-transparent background (movable/resizable)
    overlay: borderless chroma-key window rendering crisp subtitles + settings button
    """

    def __init__(self):
        global MT_CFG

        # ---------- Background window ----------
        self.bg_root = tk.Tk()
        self.bg_root.title("Live Translation (EN -> Target)")
        self.bg_root.minsize(520, 240)

        self.bg = "#101214"
        self.sep = "#2B2F36"
        self.fg = "#E8E8E8"
        self.en_fg = "#FFFFFF"
        self.out_fg = "#FFFFFF"

        self.font_family = "Segoe UI"
        self.out_font_family = "SimHei"  # default; stable
        self.font_size = 16

        self.window_alpha_var = tk.DoubleVar(value=0.5)
        self.subtitle_alpha_var = tk.DoubleVar(value=1.0)

        self.show_en_var = tk.BooleanVar(value=True)

        # One-shot action: check -> export
        self.export_history_var = tk.BooleanVar(value=False)

        # Target language selection (DISPLAY LABEL, not code)
        cur_code = (MT_CFG or {}).get("target_lang", DEFAULT_TARGET_LANG)
        default_label = NLLB_CODE_TO_LABEL.get(cur_code) or NLLB_CODE_TO_LABEL.get(DEFAULT_TARGET_LANG) or (
            NLLB_LANG_LABELS[0] if NLLB_LANG_LABELS else "Unknown language"
        )
        self.target_lang_var = tk.StringVar(value=default_label)

        # Warn user if pycountry is missing (names may become "Unknown language")
        if not HAS_PYCOUNTRY:
            try:
                self.bg_root.after(
                    200,
                    lambda: messagebox.showwarning(
                        "Language names",
                        "pycountry is not installed, so many languages may appear as 'Unknown language'.\n"
                        "Install for full English names:\n\npip install pycountry",
                    ),
                )
            except Exception:
                pass

        self.bg_root.configure(bg=self.bg)
        self.bg_root.attributes("-topmost", True)

        base_w, base_h = 520, 240
        init_w = int(base_w * 1.5)
        init_h = int(base_h * 4)
        sw = self.bg_root.winfo_screenwidth()
        sh = self.bg_root.winfo_screenheight()
        x = max(0, sw - init_w - 20)
        y = max(0, (sh - init_h) // 2)
        self.bg_root.geometry(f"{init_w}x{init_h}+{x}+{y}")

        self._apply_window_alpha()

        self.bg_fill = tk.Frame(self.bg_root, bg=self.bg, bd=0, highlightthickness=0)
        self.bg_fill.pack(fill="both", expand=True)

        # ---------- Overlay window ----------
        self.overlay = tk.Toplevel(self.bg_root)
        self.overlay.withdraw()

        self.overlay.configure(bg=KEY_COLOR)
        self.overlay.overrideredirect(True)
        self.overlay.attributes("-topmost", True)
        self._apply_subtitle_alpha()

        self._transparent_ok = False
        if platform.system().lower().startswith("win"):
            try:
                self.overlay.wm_attributes("-transparentcolor", KEY_COLOR)
                self._transparent_ok = True
            except Exception as e:
                print("[UI] transparentcolor not supported:", e)

        self.overlay_frame = tk.Frame(self.overlay, bg=KEY_COLOR, bd=0, highlightthickness=0)
        self.overlay_frame.pack(fill="both", expand=True)

        self.topbar = tk.Frame(self.overlay_frame, bg=KEY_COLOR, height=34)
        self.topbar.pack(fill="x", side="top")
        self.topbar.pack_propagate(False)

        self.settings_btn = tk.Button(
            self.topbar,
            text="⚙",
            command=self.open_settings,
            bg=KEY_COLOR,
            fg=self.fg,
            bd=0,
            highlightthickness=0,
            activebackground=KEY_COLOR,
            activeforeground=self.fg,
            cursor="hand2",
            font=(self.font_family, 11),
        )
        self.settings_btn.pack(side="right", padx=10, pady=4, ipadx=0, ipady=0)

        self.content = tk.Frame(self.overlay_frame, bg=KEY_COLOR, bd=0, highlightthickness=0)
        self.content.pack(fill="both", expand=True)

        self.history_text = tk.Text(
            self.content,
            bg=KEY_COLOR,
            fg=self.fg,
            bd=0,
            highlightthickness=0,
            wrap="word",
            font=(self.font_family, self.font_size),
            height=16,
        )
        self.history_text.pack(fill="both", expand=True, padx=10, pady=(6, 6))

        self.sep_line = tk.Frame(self.content, bg=self.sep, height=2)
        self.sep_line.pack(fill="x", padx=10, pady=0)

        self.current_text = tk.Text(
            self.content,
            bg=KEY_COLOR,
            fg=self.fg,
            bd=0,
            highlightthickness=0,
            wrap="word",
            font=(self.font_family, self.font_size),
            height=5,
        )
        self.current_text.pack(fill="x", padx=10, pady=(6, 10))

        self.history_text.configure(state="disabled")
        self.current_text.configure(state="disabled")

        self._apply_subtitle_colors()

        self.menu = tk.Menu(self.overlay, tearoff=0)
        self.menu.add_command(label="Choose EN Color", command=self.choose_en_color)
        self.menu.add_command(label="Choose OUT Color", command=self.choose_out_color)
        self.menu.add_separator()
        self.menu.add_command(label="Quit", command=self.on_close)

        self.overlay.bind("<Button-3>", self.show_menu)
        self.bg_root.bind("<Button-3>", self.show_menu)

        self.bg_root.bind("<Configure>", lambda _e: self.sync_overlay_geometry())
        self.bg_root.bind("<Unmap>", lambda _e: self.overlay.withdraw())
        self.bg_root.bind("<Map>", lambda _e: self._on_bg_map())

        self.bg_root.protocol("WM_DELETE_WINDOW", self.on_close)

        self._export_progress_q = queue.Queue()
        self._exporting = False

        self.bg_root.after(80, self._initial_show)
        self.bg_root.after(int(POLL_INTERVAL * 1000), self.tick_ui)

    def _initial_show(self):
        self.sync_overlay_geometry()
        try:
            self.overlay.deiconify()
            self.overlay.lift()
            self.settings_btn.lift()
        except Exception:
            pass

    def _on_bg_map(self):
        self.sync_overlay_geometry()
        try:
            self.overlay.deiconify()
            self.overlay.lift()
            self.settings_btn.lift()
        except Exception:
            pass

    def _apply_window_alpha(self):
        try:
            a = float(self.window_alpha_var.get())
            a = max(0.02, min(1.0, a))
            self.bg_root.attributes("-alpha", a)
        except Exception:
            pass

    def _apply_subtitle_alpha(self):
        try:
            a = float(self.subtitle_alpha_var.get())
            a = max(0.05, min(1.0, a))
            self.overlay.attributes("-alpha", a)
        except Exception:
            pass

    def _apply_subtitle_colors(self):
        self.history_text.tag_configure("en", foreground=self.en_fg, font=(self.font_family, self.font_size))
        self.history_text.tag_configure("out", foreground=self.out_fg, font=(self.out_font_family, self.font_size))
        self.current_text.tag_configure("en", foreground=self.en_fg, font=(self.font_family, self.font_size))
        self.current_text.tag_configure("out", foreground=self.out_fg, font=(self.out_font_family, self.font_size))

    def choose_en_color(self):
        c = colorchooser.askcolor(title="EN Color", initialcolor=self.en_fg)[1]
        if c:
            self.en_fg = c
            self._apply_subtitle_colors()

    def choose_out_color(self):
        c = colorchooser.askcolor(title="Output Color", initialcolor=self.out_fg)[1]
        if c:
            self.out_fg = c
            self._apply_subtitle_colors()

    def choose_window_bg(self):
        c = colorchooser.askcolor(title="Window Background", initialcolor=self.bg)[1]
        if c:
            self.bg = c
            self.bg_root.configure(bg=self.bg)
            self.bg_fill.configure(bg=self.bg)

    def sync_overlay_geometry(self):
        try:
            x = self.bg_root.winfo_rootx()
            y = self.bg_root.winfo_rooty()
            w = self.bg_root.winfo_width()
            h = self.bg_root.winfo_height()
            if w < 50 or h < 50:
                return
            self.overlay.geometry(f"{w}x{h}+{x}+{y}")
            self.overlay.lift()
            self.settings_btn.lift()
        except Exception:
            pass

    def show_menu(self, event):
        try:
            self.menu.tk_popup(event.x_root, event.y_root)
        finally:
            try:
                self.menu.grab_release()
            except Exception:
                pass

    # ================= Export (Word from SQLite, auto-split) =================

    def _make_export_modal(self):
        win = tk.Toplevel(self.bg_root)
        win.title("Exporting...")
        win.attributes("-topmost", True)
        win.resizable(False, False)
        win.configure(bg=self.bg)

        lbl = tk.Label(win, text="Exporting THIS SESSION history to Word...\nPlease wait.", bg=self.bg, fg=self.fg)
        lbl.pack(padx=16, pady=(16, 8))

        self._export_status_var = tk.StringVar(value="Preparing...")
        status = tk.Label(win, textvariable=self._export_status_var, bg=self.bg, fg=self.fg)
        status.pack(padx=16, pady=(0, 16))

        win.update_idletasks()
        sw = win.winfo_screenwidth()
        sh = win.winfo_screenheight()
        ww = win.winfo_width()
        wh = win.winfo_height()
        win.geometry(f"{ww}x{wh}+{(sw - ww) // 2}+{(sh - wh) // 2}")
        return win

    def _export_to_docx_parts(self, base_path: str) -> List[str]:
        """
        Export persisted records for the CURRENT SESSION (SESSION_ID) from SQLite.
        Returns list of written file paths.
        """
        assert HISTORY_DB_PATH is not None
        try:
            from docx import Document
        except Exception as e:
            raise RuntimeError(f"python-docx not available: {e}")

        conn = _db_connect(HISTORY_DB_PATH)
        try:
            total, max_id = _fetch_session_bounds(conn, SESSION_ID)
            if total <= 0 or max_id <= 0:
                return []

            lang_summary = _fetch_session_lang_summary(conn, SESSION_ID)

            base_dir = os.path.dirname(os.path.abspath(base_path))
            base_name = os.path.splitext(os.path.basename(base_path))[0]
            parts = (total + DOCX_MAX_RECORDS_PER_FILE - 1) // DOCX_MAX_RECORDS_PER_FILE
            written: List[str] = []

            def _start_doc(part_idx: int):
                d = Document()
                d.add_heading("Translation History (EN -> Target)", level=1)
                d.add_paragraph("Scope: CURRENT SESSION only (this run).")
                d.add_paragraph(f"Session: {SESSION_ID}")
                d.add_paragraph(f"Target language(s): {lang_summary}")
                d.add_paragraph(time.strftime("Exported at: %Y-%m-%d %H:%M:%S"))
                if parts > 1:
                    d.add_paragraph(f"Part {part_idx} of {parts} (max {DOCX_MAX_RECORDS_PER_FILE} records per file)")
                d.add_paragraph("")
                return d

            current_part = 1
            current_in_part = 0
            doc = _start_doc(current_part)

            record_index = 0
            for (_id, ts_str, tgt_lang, en, out_text) in iter_session_records(conn, SESSION_ID, max_id, batch_size=500):
                record_index += 1
                tgt_lang = (tgt_lang or "").strip() or DEFAULT_TARGET_LANG

                doc.add_paragraph(f"{record_index}.  [{ts_str}]  ({tgt_lang})")
                if en:
                    doc.add_paragraph(f"EN: {en}")
                if out_text:
                    doc.add_paragraph(f"OUT: {out_text}")
                doc.add_paragraph("")

                current_in_part += 1
                if current_in_part >= DOCX_MAX_RECORDS_PER_FILE:
                    if parts == 1:
                        out_path = os.path.join(base_dir, f"{base_name}.docx")
                    else:
                        out_path = os.path.join(base_dir, f"{base_name}_part{current_part:03d}.docx")
                    doc.save(out_path)
                    written.append(out_path)

                    current_part += 1
                    current_in_part = 0
                    if current_part <= parts:
                        doc = _start_doc(current_part)

            if current_in_part > 0:
                if parts == 1:
                    out_path = os.path.join(base_dir, f"{base_name}.docx")
                else:
                    out_path = os.path.join(base_dir, f"{base_name}_part{current_part:03d}.docx")
                doc.save(out_path)
                written.append(out_path)

            return written
        finally:
            conn.close()

    def _export_worker(self, base_path: str):
        try:
            self._export_progress_q.put(("status", "Reading persisted history..."))
            paths = self._export_to_docx_parts(base_path)
            if not paths:
                self._export_progress_q.put(("done", ("empty", [])))
                return
            self._export_progress_q.put(("done", ("ok", paths)))
        except Exception as e:
            self._export_progress_q.put(("done", ("err", str(e))))

    def _poll_export_progress(self):
        if not self._exporting:
            return

        try:
            while True:
                kind, payload = self._export_progress_q.get_nowait()
                if kind == "status":
                    try:
                        self._export_status_var.set(str(payload))
                    except Exception:
                        pass
                elif kind == "done":
                    status, data = payload
                    self._exporting = False
                    try:
                        if getattr(self, "_export_modal", None) is not None and self._export_modal.winfo_exists():
                            self._export_modal.destroy()
                    except Exception:
                        pass

                    if status == "empty":
                        messagebox.showinfo("Export", "No persisted history found for this session yet.")
                    elif status == "ok":
                        paths = data
                        if len(paths) == 1:
                            messagebox.showinfo("Export", f"Exported successfully:\n{paths[0]}")
                        else:
                            messagebox.showinfo(
                                "Export",
                                "Exported successfully (auto-split into multiple files):\n\n"
                                + "\n".join(paths[:8])
                                + ("" if len(paths) <= 8 else f"\n... (+{len(paths) - 8} more)"),
                            )
                    else:
                        messagebox.showerror("Export failed", str(data))
                    return
        except queue.Empty:
            pass

        self.bg_root.after(120, self._poll_export_progress)

    def _on_export_history_toggle(self):
        if not self.export_history_var.get():
            return

        try:
            default_name = f"translation_history_{SESSION_ID}.docx"
            path = filedialog.asksaveasfilename(
                parent=getattr(self, "settings_win", None),
                title="Export THIS SESSION history (Word .docx, auto-split if large)",
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                initialfile=default_name,
            )
            if not path:
                return
            if not path.lower().endswith(".docx"):
                path = path + ".docx"

            if self._exporting:
                messagebox.showinfo("Export", "An export is already running.")
                return

            self._exporting = True
            self._export_modal = self._make_export_modal()
            self._export_status_var.set("Starting export...")

            t = threading.Thread(target=self._export_worker, args=(path,), daemon=True)
            t.start()

            self.bg_root.after(120, self._poll_export_progress)

        finally:
            self.export_history_var.set(False)

    # ================= UI loop and close =================

    def on_close(self):
        global RUNNING
        RUNNING = False
        try:
            stop_audio_stream()
        except Exception:
            pass
        try:
            if getattr(self, "settings_win", None) is not None and self.settings_win.winfo_exists():
                self.settings_win.destroy()
        except Exception:
            pass
        try:
            if getattr(self, "_export_modal", None) is not None and self._export_modal.winfo_exists():
                self._export_modal.destroy()
        except Exception:
            pass
        try:
            self.overlay.destroy()
        except Exception:
            pass
        try:
            self.bg_root.destroy()
        except Exception:
            pass

    def render_history(self):
        self.history_text.configure(state="normal")
        self.history_text.delete("1.0", "end")
        for en, out_text in history_segments[-MAX_HISTORY_SEGMENTS:]:
            if self.show_en_var.get() and en:
                self.history_text.insert("end", en.strip() + "\n", ("en",))
            if out_text:
                self.history_text.insert("end", out_text.strip() + "\n", ("out",))
            self.history_text.insert("end", "\n")
        self.history_text.configure(state="disabled")
        self.history_text.see("end")

    def render_current(self):
        self.current_text.configure(state="normal")
        self.current_text.delete("1.0", "end")
        if self.show_en_var.get() and current_en:
            self.current_text.insert("end", current_en.strip() + "\n", ("en",))
        if current_out:
            self.current_text.insert("end", current_out.strip() + "\n", ("out",))
        self.current_text.configure(state="disabled")

    def tick_ui(self):
        self.render_history()
        self.render_current()
        self.sync_overlay_geometry()
        try:
            self.settings_btn.lift()
        except Exception:
            pass
        if RUNNING:
            self.bg_root.after(int(POLL_INTERVAL * 1000), self.tick_ui)

    # ================= Settings =================

    def _device_choices(self):
        choices = ["Default (system)"]
        mapping = {"Default (system)": None}
        for idx, name in list_input_devices():
            label = f"[{idx}] {name}"
            choices.append(label)
            mapping[label] = idx
        return choices, mapping

    def open_settings(self):
        if getattr(self, "settings_win", None) is not None and self.settings_win.winfo_exists():
            self.settings_win.lift()
            return

        win = tk.Toplevel(self.bg_root)
        self.settings_win = win
        win.title("Settings")
        win.attributes("-topmost", True)
        win.configure(bg=self.bg)

        pad = 10
        frm = tk.Frame(win, bg=self.bg)
        frm.pack(fill="both", expand=True, padx=pad, pady=pad)

        row = 0

        chk = tk.Checkbutton(
            frm,
            text="Show English subtitles",
            variable=self.show_en_var,
            bg=self.bg, fg=self.fg,
            activebackground=self.bg, activeforeground=self.fg,
            selectcolor=self.bg,
        )
        chk.grid(row=row, column=0, columnspan=2, sticky="w")
        row += 1

        # Target language dropdown (ENGLISH NAME labels)
        tk.Label(frm, text="Target language", bg=self.bg, fg=self.fg).grid(
            row=row, column=0, sticky="w", pady=(10, 0)
        )

        cb_lang = ttk.Combobox(
            frm,
            textvariable=self.target_lang_var,
            values=NLLB_LANG_LABELS,
            state="readonly",
            width=38,
        )
        cb_lang.grid(row=row, column=1, sticky="e", pady=(10, 0))

        def _apply_target_lang(_evt=None):
            label = (self.target_lang_var.get() or "").strip()
            code = NLLB_LABEL_TO_CODE.get(label, "")
            if not code:
                messagebox.showerror("Target language", "Invalid selection.")
                return
            ok = set_target_language(code)
            if not ok:
                messagebox.showerror("Target language", "Failed to apply target language.")

        cb_lang.bind("<<ComboboxSelected>>", _apply_target_lang)
        row += 1

        tk.Label(frm, text="English color", bg=self.bg, fg=self.fg).grid(row=row, column=0, sticky="w", pady=(10, 0))
        tk.Button(frm, text="Pick", command=self.choose_en_color).grid(row=row, column=1, sticky="e", pady=(10, 0))
        row += 1

        tk.Label(frm, text="Output color", bg=self.bg, fg=self.fg).grid(row=row, column=0, sticky="w")
        tk.Button(frm, text="Pick", command=self.choose_out_color).grid(row=row, column=1, sticky="e")
        row += 1

        tk.Label(frm, text="Window background", bg=self.bg, fg=self.fg).grid(row=row, column=0, sticky="w", pady=(10, 0))
        tk.Button(frm, text="Pick", command=self.choose_window_bg).grid(row=row, column=1, sticky="e", pady=(10, 0))
        row += 1

        tk.Label(frm, text="Window transparency (background)", bg=self.bg, fg=self.fg).grid(row=row, column=0, sticky="w")
        s1 = tk.Scale(
            frm, from_=0.02, to=1.0,
            resolution=0.01,
            orient="horizontal",
            variable=self.window_alpha_var,
            command=lambda _v: self._apply_window_alpha(),
            bg=self.bg, fg=self.fg,
            highlightthickness=0,
            troughcolor=self.sep,
            length=200,
        )
        s1.grid(row=row, column=1, sticky="e")
        row += 1

        tk.Label(frm, text="Subtitle opacity (EN+OUT)", bg=self.bg, fg=self.fg).grid(row=row, column=0, sticky="w", pady=(10, 0))
        s2 = tk.Scale(
            frm, from_=0.05, to=1.0,
            resolution=0.01,
            orient="horizontal",
            variable=self.subtitle_alpha_var,
            command=lambda _v: self._apply_subtitle_alpha(),
            bg=self.bg, fg=self.fg,
            highlightthickness=0,
            troughcolor=self.sep,
            length=200,
        )
        s2.grid(row=row, column=1, sticky="e", pady=(10, 0))
        row += 1

        tk.Label(frm, text="Audio input (select to apply)", bg=self.bg, fg=self.fg).grid(row=row, column=0, sticky="w", pady=(10, 0))

        choices, mapping = self._device_choices()
        self._device_map = mapping

        cur_label = "Default (system)"
        for k, v in mapping.items():
            if v == AUDIO_DEVICE_INDEX:
                cur_label = k
                break

        self.device_var = tk.StringVar(value=cur_label)
        cb_dev = ttk.Combobox(frm, textvariable=self.device_var, values=choices, state="readonly", width=30)
        cb_dev.grid(row=row, column=1, sticky="e", pady=(10, 0))

        def _apply_device_from_combo(_evt=None):
            label = self.device_var.get()
            idx = self._device_map.get(label, None)
            set_audio_device(idx)

        cb_dev.bind("<<ComboboxSelected>>", _apply_device_from_combo)
        row += 1

        export_chk = tk.Checkbutton(
            frm,
            text=f"Export THIS SESSION history (Word .docx, auto-split every {DOCX_MAX_RECORDS_PER_FILE} records)",
            variable=self.export_history_var,
            command=self._on_export_history_toggle,
            bg=self.bg, fg=self.fg,
            activebackground=self.bg, activeforeground=self.fg,
            selectcolor=self.bg,
        )
        export_chk.grid(row=row, column=0, columnspan=2, sticky="w", pady=(12, 0))
        row += 1

        tip = "OK (Windows)" if self._transparent_ok else "NOT SUPPORTED (likely non-Windows)"
        tk.Label(frm, text=f"Transparent overlay: {tip}", bg=self.bg, fg=self.fg).grid(
            row=row, column=0, columnspan=2, sticky="w", pady=(12, 0)
        )
        row += 1

        tk.Label(frm, text=f"History DB: {HISTORY_DB_PATH}", bg=self.bg, fg=self.fg).grid(
            row=row, column=0, columnspan=2, sticky="w", pady=(8, 0)
        )
        row += 1

        tk.Label(frm, text=f"Session: {SESSION_ID}", bg=self.bg, fg=self.fg).grid(
            row=row, column=0, columnspan=2, sticky="w", pady=(2, 0)
        )
        row += 1

        frm.grid_columnconfigure(0, weight=1)
        frm.grid_columnconfigure(1, weight=0)

    def mainloop(self):
        self.bg_root.mainloop()


# ===================== Worker Thread: ASR + Translation + Finalize =====================

def _commit_history(en: str, out_text: str) -> None:
    """
    Commit (en, out_text) to:
      1) UI ring buffer
      2) SQLite (async queue, best-effort)
    """
    global history_segments

    en = (en or "").strip()
    out_text = (out_text or "").strip()
    if not en and not out_text:
        return

    history_segments.append((en, out_text))
    if len(history_segments) > MAX_HISTORY_SEGMENTS:
        history_segments.pop(0)

    enqueue_history_record(en, out_text)


def asr_worker():
    global audio_buffer, pending_en, current_en, current_out, asr_last_text_ts

    if not load_whisper_model():
        print("[ASR] Whisper not available.")
        return
    if not start_audio_stream():
        print("[ASR] Audio stream not available.")
        return

    min_samples = int(AUDIO_SAMPLE_RATE * WHISPER_MIN_SECONDS)
    target_samples = int(AUDIO_SAMPLE_RATE * WHISPER_CHUNK_SECONDS)

    overlap_samples = int(AUDIO_SAMPLE_RATE * AUDIO_OVERLAP_SECONDS)
    prev_tail = np.zeros(0, dtype=np.float32)

    last_full_text = ""

    while RUNNING:
        try:
            chunk = audio_q.get(timeout=0.2)
            if chunk is not None and chunk.size > 0:
                audio_buffer = np.concatenate([audio_buffer, chunk])
        except queue.Empty:
            pass

        now = time.time()

        if audio_buffer.size >= min_samples:
            take = min(audio_buffer.size, target_samples)
            cur = audio_buffer[:take].copy()
            audio_buffer = audio_buffer[take:]

            had_overlap = (overlap_samples > 0 and prev_tail.size > 0)

            # Prepend previous tail as overlap prefix to reduce boundary word drops.
            audio = np.concatenate([prev_tail, cur]) if had_overlap else cur
            audio = normalize_audio(audio)

            # Update tail to carry into next chunk.
            if overlap_samples > 0 and cur.size > 0:
                keep = min(cur.size, overlap_samples)
                prev_tail = cur[-keep:].copy()
            else:
                prev_tail = np.zeros(0, dtype=np.float32)

            try:
                segments, _info = whisper_model.transcribe(
                    audio,
                    language="en",
                    beam_size=5,
                    vad_filter=True,
                )

                texts = []
                for seg in segments:
                    t = (getattr(seg, "text", "") or "").strip()
                    if not t:
                        continue

                    if had_overlap:
                        # Drop segments likely from overlap prefix window (heuristic).
                        try:
                            seg_end = getattr(seg, "end", None)
                            if seg_end is not None and seg_end <= AUDIO_OVERLAP_SECONDS * 0.95:
                                continue
                        except Exception:
                            pass

                    texts.append(t)

                full_text = re.sub(r"\s+", " ", " ".join(texts)).strip()

            except Exception as e:
                print(f"[ASR] transcribe error: {e}")
                continue

            if full_text and is_valid_en_segment(full_text):
                increment = full_text

                # Remove duplicated prefix due to overlap between consecutive ASR runs.
                if last_full_text:
                    k_words = word_overlap_suffix_prefix(last_full_text, full_text, max_words=8)
                    if k_words >= 2:
                        increment = drop_prefix_words(full_text, k_words)
                    else:
                        k = overlap_suffix_prefix(last_full_text, full_text, max_k=MAX_OVERLAP_CHARS)
                        if k >= 8:
                            increment = full_text[k:].lstrip(" \t\r\n,.;:!?")

                last_full_text = full_text

                if increment and is_valid_en_segment(increment):
                    asr_last_text_ts = now
                    pending_en = append_with_overlap(pending_en, increment)

                    completed, remainder = split_complete_sentences(pending_en)

                    # Keep the last sentence as "current line" unless we are sure it's finished.
                    to_commit = completed
                    if completed and (remainder or "") == "":
                        to_commit = completed[:-1]
                        remainder = completed[-1]

                    if to_commit:
                        for s in to_commit:
                            s_norm = s.strip()
                            if not s_norm:
                                continue
                            if history_segments and history_segments[-1][0].strip().lower() == s_norm.lower():
                                continue

                            out_text = translate_en_to_target(s_norm)
                            _commit_history(s_norm, out_text)

                    pending_en = remainder
                    current_en = pending_en
                    current_out = translate_en_to_target(current_en) if current_en else ""

        # Idle finalize: commit the current line after silence
        if pending_en and (now - asr_last_text_ts) >= IDLE_FINALIZE_SECONDS:
            if len(pending_en.strip()) >= MIN_FINALIZE_CHARS:
                s = pending_en.strip()
                if not history_segments or history_segments[-1][0].strip().lower() != s.lower():
                    out_text = translate_en_to_target(s)
                    _commit_history(s, out_text)

                pending_en = ""
                current_en = ""
                current_out = ""

        time.sleep(0.01)


# ===================== Entry Point =====================

def _parse_args_apply_env():
    parser = argparse.ArgumentParser(description="Live Translation Overlay (Whisper ASR + NLLB MT)")
    parser.add_argument("--hf-home", default=None, help="Override HF_HOME (Hugging Face cache dir)")
    parser.add_argument("--db-path", default=None, help="Override HISTORY_DB_PATH (SQLite file path)")
    parser.add_argument(
        "--disable-hf-transfer",
        action="store_true",
        help="Disable HF_HUB_ENABLE_HF_TRANSFER (sometimes helps if downloads fail).",
    )
    return parser.parse_args()


def main():
    global RUNNING, HISTORY_DB_PATH

    args = _parse_args_apply_env()

    if args.hf_home:
        os.environ["HF_HOME"] = args.hf_home
    if args.db_path:
        os.environ["HISTORY_DB_PATH"] = args.db_path
    if args.disable_hf_transfer:
        os.environ["HF_HUB_ENABLE_HF_TRANSFER"] = "0"

    HISTORY_DB_PATH = os.environ.get("HISTORY_DB_PATH", DEFAULT_DB_PATH)

    # Defer MT load until after env/CLI overrides
    bootstrap_mt_and_languages()

    start_logger()

    app = OverlaySubtitleApp()

    t = threading.Thread(target=asr_worker, daemon=True)
    t.start()

    try:
        app.mainloop()
    finally:
        RUNNING = False
        try:
            stop_audio_stream()
        except Exception:
            pass
        stop_logger()


if __name__ == "__main__":
    main()
